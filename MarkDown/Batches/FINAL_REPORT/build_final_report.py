import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

repo_root = r"C:\Users\chloud\.gemini\antigravity\worktrees\justificadll\draft_final_report_justifiqa"
base_dir = os.path.join(repo_root, "MarkDown", "FinalReport")
batch_dir = os.path.join(repo_root, "MarkDown", "Batches", "FINAL_REPORT")

files = [
    "00_FRONT_MATTER.md",
    "01_BAB_I_PENDAHULUAN.md",
    "02_BAB_II_LANDASAN_TEORI.md",
    "03_BAB_III_ANALISIS_DAN_PERANCANGAN.md",
    "04_BAB_IV_IMPLEMENTASI_DAN_PENGUJIAN.md",
    "05_BAB_V_PENUTUP.md",
    "06_DAFTAR_PUSTAKA.md",
    "07_LAMPIRAN.md"
]

def strip_trailing_whitespace(text):
    lines = text.splitlines()
    cleaned_lines = [line.rstrip() for line in lines]
    return "\n".join(cleaned_lines) + "\n"

# Clean trailing whitespace across all report markdown files
for filename in files:
    filepath = os.path.join(base_dir, filename)
    if os.path.exists(filepath):
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        cleaned = strip_trailing_whitespace(content)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(cleaned)

for filename in ["FIGURE_PLAN.md", "SOURCE_TRACEABILITY.md"]:
    filepath = os.path.join(base_dir, filename)
    if os.path.exists(filepath):
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        cleaned = strip_trailing_whitespace(content)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(cleaned)

for filename in ["BATCH.md", "PROMPT_MASTER.md", "LEARNING.md"]:
    filepath = os.path.join(batch_dir, filename)
    if os.path.exists(filepath):
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
        cleaned = strip_trailing_whitespace(content)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(cleaned)

# 1. Build combined LAPORAN_TUGAS_AKHIR_JUSTIFIQA.md
combined_md_path = os.path.join(base_dir, "LAPORAN_TUGAS_AKHIR_JUSTIFIQA.md")
combined_content = []

for filename in files:
    filepath = os.path.join(base_dir, filename)
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
        combined_content.append(content.strip())

full_md_text = "\n\n---\n\n".join(combined_content) + "\n"

with open(combined_md_path, "w", encoding="utf-8") as f:
    f.write(full_md_text)

print(f"Combined Markdown written to {combined_md_path} ({len(full_md_text)} bytes)")

# 2. Build LAPORAN_TUGAS_AKHIR_JUSTIFIQA.docx
docx_path = os.path.join(base_dir, "LAPORAN_TUGAS_AKHIR_JUSTIFIQA.docx")
doc = Document()

# Page Setup: A4, 1.0 inch margins all around
for section in doc.sections:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Footer PAGE field (wrapped inside <w:r> runs for OOXML compliance)
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Halaman ")
    footer_run.font.name = 'Times New Roman'
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(0, 0, 0)
    
    r_begin = footer_p.add_run()
    r_begin._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))

    r_instr = footer_p.add_run()
    r_instr._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))

    r_sep = footer_p.add_run()
    r_sep._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))

    r_end = footer_p.add_run()
    r_end._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

# Styles Setup: 100% Black Text Only
normal_style = doc.styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(12)
normal_style.font.color.rgb = RGBColor(0, 0, 0)
normal_style.paragraph_format.line_spacing = 1.5
normal_style.paragraph_format.space_after = Pt(6)

try:
    title_style = doc.styles['Title']
except KeyError:
    title_style = doc.styles.add_style('Title', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
title_style.font.name = 'Times New Roman'
title_style.font.size = Pt(18)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0, 0, 0)
title_style.paragraph_format.space_before = Pt(12)
title_style.paragraph_format.space_after = Pt(12)
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_style.paragraph_format.first_line_indent = Inches(0)

# Heading 1: Centered, Bold, Black, 14pt, 0 first line indent
h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Times New Roman'
h1_style.font.size = Pt(14)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0, 0, 0)
h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1_style.paragraph_format.space_before = Pt(14)
h1_style.paragraph_format.space_after = Pt(6)
h1_style.paragraph_format.line_spacing = 1.15
h1_style.paragraph_format.keep_with_next = True
h1_style.paragraph_format.first_line_indent = Inches(0)

# Heading 2: Left Aligned, Bold, Black, 12pt, 0 first line indent
h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Times New Roman'
h2_style.font.size = Pt(12)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0, 0, 0)
h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2_style.paragraph_format.space_before = Pt(12)
h2_style.paragraph_format.space_after = Pt(6)
h2_style.paragraph_format.line_spacing = 1.15
h2_style.paragraph_format.keep_with_next = True
h2_style.paragraph_format.first_line_indent = Inches(0)

# Heading 3: Left Aligned, Bold Italic, Black, 12pt, 0 first line indent
h3_style = doc.styles['Heading 3']
h3_style.font.name = 'Times New Roman'
h3_style.font.size = Pt(12)
h3_style.font.bold = True
h3_style.font.italic = True
h3_style.font.color.rgb = RGBColor(0, 0, 0)
h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h3_style.paragraph_format.space_before = Pt(10)
h3_style.paragraph_format.space_after = Pt(4)
h3_style.paragraph_format.line_spacing = 1.15
h3_style.paragraph_format.keep_with_next = True
h3_style.paragraph_format.first_line_indent = Inches(0)

try:
    h4_style = doc.styles['Heading 4']
except KeyError:
    h4_style = doc.styles.add_style('Heading 4', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
h4_style.font.name = 'Times New Roman'
h4_style.font.size = Pt(12)
h4_style.font.italic = True
h4_style.font.color.rgb = RGBColor(0, 0, 0)
h4_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h4_style.paragraph_format.space_before = Pt(8)
h4_style.paragraph_format.space_after = Pt(4)
h4_style.paragraph_format.line_spacing = 1.15
h4_style.paragraph_format.keep_with_next = True
h4_style.paragraph_format.first_line_indent = Inches(0)

def parse_inline_markdown(paragraph, text, font_name='Times New Roman', font_size=Pt(12), text_color=RGBColor(0, 0, 0), default_bold=False, default_italic=False):
    """
    Parses inline markdown markers (**bold**, *italic*, `code`) into Word runs.
    Code spans (`...`) preserve literal characters inside without italicizing wildcards.
    """
    code_pattern = re.compile(r'(`[^`]+`)')
    parts = code_pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('`') and part.endswith('`'):
            code_text = part[1:-1]
            run = paragraph.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(40, 40, 40)
            run.bold = default_bold
            run.italic = default_italic
        else:
            b_i_pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)')
            subparts = b_i_pattern.split(part)
            for sub in subparts:
                if not sub:
                    continue
                run = paragraph.add_run()
                run.font.name = font_name
                run.font.size = font_size
                run.font.color.rgb = text_color
                
                if sub.startswith('***') and sub.endswith('***'):
                    run.text = sub[3:-3]
                    run.bold = True
                    run.italic = True
                elif sub.startswith('**') and sub.endswith('**'):
                    run.text = sub[2:-2]
                    run.bold = True
                    run.italic = default_italic
                elif sub.startswith('*') and sub.endswith('*'):
                    run.text = sub[1:-1]
                    run.bold = default_bold
                    run.italic = True
                else:
                    run.text = sub
                    run.bold = default_bold
                    run.italic = default_italic

def add_table_styled(doc, rows_data):
    if not rows_data:
        return
    num_rows = len(rows_data)
    num_cols = max(len(r) for r in rows_data)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Repeat header row across pages
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    
    # Set clean grayscale borders
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="CCCCCC"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)
    
    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = "" # Clear default text
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.first_line_indent = Inches(0)
                
                is_header = (r_idx == 0)
                parse_inline_markdown(p, val, font_name='Times New Roman', font_size=Pt(10), text_color=RGBColor(0, 0, 0), default_bold=is_header)
                
                if is_header:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F2F2F2"/>')
                    cell._tc.get_or_add_tcPr().append(shading)

    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_after = Pt(6)
    p_after.paragraph_format.first_line_indent = Inches(0)

def process_markdown_file_to_docx(doc, file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        lines = f.readlines()
    
    in_table = False
    table_lines = []
    in_codeblock = False
    codeblock_lines = []

    def flush_table_buffer(doc, t_lines):
        if not t_lines:
            return
        rows_data = []
        for tl in t_lines:
            tl = tl.strip()
            if not tl.startswith('|'):
                continue
            parts = [p.strip() for p in tl.split('|')[1:-1]]
            if all(re.match(r'^:?-+:?$', p) for p in parts if p):
                continue
            rows_data.append(parts)
        if rows_data:
            add_table_styled(doc, rows_data)

    def flush_codeblock(doc, c_lines):
        if not c_lines:
            return
        code_text = "".join(c_lines)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(0)
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(40, 40, 40)

    for line in lines:
        stripped = line.strip()
        
        # Codeblock boundary
        if stripped.startswith('```'):
            if in_codeblock:
                in_codeblock = False
                flush_codeblock(doc, codeblock_lines)
                codeblock_lines = []
            else:
                if in_table:
                    in_table = False
                    flush_table_buffer(doc, table_lines)
                    table_lines = []
                in_codeblock = True
            continue
        
        if in_codeblock:
            codeblock_lines.append(line)
            continue
        
        # Table row boundary
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_lines = [stripped]
            else:
                table_lines.append(stripped)
            continue
        else:
            if in_table:
                in_table = False
                flush_table_buffer(doc, table_lines)
                table_lines = []

        if not stripped:
            continue

        # Document Title
        if stripped.startswith('# LAPORAN TUGAS AKHIR'):
            p = doc.add_paragraph(style='Title')
            p.paragraph_format.first_line_indent = Inches(0)
            parse_inline_markdown(p, stripped[2:], font_size=Pt(18), default_bold=True)
        # Heading 1 (BAB headings)
        elif stripped.startswith('# '):
            p = doc.add_paragraph(style='Heading 1')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Inches(0)
            heading_text = stripped[2:]
            parse_inline_markdown(p, heading_text, font_size=Pt(14), default_bold=True)
        # Heading 2 (Numbered subchapters)
        elif stripped.startswith('## '):
            p = doc.add_paragraph(style='Heading 2')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0)
            parse_inline_markdown(p, stripped[3:], font_size=Pt(12), default_bold=True)
        # Heading 3 (Lower subsections)
        elif stripped.startswith('### '):
            p = doc.add_paragraph(style='Heading 3')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0)
            parse_inline_markdown(p, stripped[4:], font_size=Pt(12), default_bold=True, default_italic=True)
        # Heading 4
        elif stripped.startswith('#### '):
            p = doc.add_paragraph(style='Heading 4')
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Inches(0)
            parse_inline_markdown(p, stripped[5:], font_size=Pt(12), default_italic=True)
        # Horizontal Rule
        elif stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run("_________________________________________________________________________________")
            run.font.color.rgb = RGBColor(200, 200, 200)
        # Captions (Gambar ..., Tabel ...)
        elif re.match(r'^\*?(Gambar|Tabel)\s+\d+', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.keep_with_next = True
            parse_inline_markdown(p, stripped, default_italic=True)
        # Bullet Lists (- or *)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25) # Hanging indent
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            bullet_run = p.add_run("• ")
            bullet_run.font.name = 'Times New Roman'
            bullet_run.font.size = Pt(12)
            bullet_run.font.color.rgb = RGBColor(0, 0, 0)
            
            parse_inline_markdown(p, stripped[2:])
        # Numbered Lists (1. , 2. , etc.) - Deterministic Isolated Hanging Indent
        elif re.match(r'^\d+\.\s', stripped):
            num_match = re.match(r'^(\d+\.)\s+(.*)', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25) # Hanging indent
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            if num_match:
                prefix = num_match.group(1) + " "
                body = num_match.group(2)
                num_run = p.add_run(prefix)
                num_run.font.name = 'Times New Roman'
                num_run.font.size = Pt(12)
                num_run.font.bold = True
                num_run.font.color.rgb = RGBColor(0, 0, 0)
                
                parse_inline_markdown(p, body)
            else:
                parse_inline_markdown(p, stripped)
        # Blockquotes (> ...)
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(0)
            parse_inline_markdown(p, stripped[2:], default_italic=True)
        # Normal Body Paragraphs (Justified, 1.5 line spacing, 1.27 cm first line indent)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.first_line_indent = Inches(0.5) # 1.27 cm first-line indent
            parse_inline_markdown(p, stripped)

    # CRITICAL FIX: Flush table if file ends while in_table is True!
    if in_table:
        flush_table_buffer(doc, table_lines)
    if in_codeblock:
        flush_codeblock(doc, codeblock_lines)

for idx, fname in enumerate(files):
    fpath = os.path.join(base_dir, fname)
    process_markdown_file_to_docx(doc, fpath)
    if idx < len(files) - 1:
        doc.add_page_break()

doc.save(docx_path)
print(f"DOCX report successfully generated at {docx_path}")
